import csv
import io
import json
from pathlib import Path

import fitz
from docx import Document
from docx.shared import Pt, Inches
from PIL import Image

from app.models import ImageFormat


class PDFConverter:
    @staticmethod
    def to_docx(input_path: str, output_path: str, progress_callback=None) -> str:
        pdf_document = fitz.open(input_path)
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)

        total_pages = len(pdf_document)

        for page_num in range(total_pages):
            page = pdf_document[page_num]
            blocks = page.get_text("blocks")

            for block in blocks:
                if block[6] == 0:
                    text = block[4].strip()
                    if text:
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(text)
                        run.font.size = Pt(11)

            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                try:
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]

                    img_path = Path(output_path).parent / f"temp_img_{page_num}_{img_index}.png"
                    with open(img_path, "wb") as img_file:
                        img_file.write(image_bytes)

                    doc.add_picture(str(img_path), width=Inches(5))

                    img_path.unlink(missing_ok=True)
                except Exception:
                    pass

            if page_num < total_pages - 1:
                doc.add_page_break()

            if progress_callback:
                progress = int(((page_num + 1) / total_pages) * 100)
                progress_callback(progress)

        pdf_document.close()
        doc.save(output_path)
        return output_path


class ImageConverter:
    FORMAT_MAP = {
        ImageFormat.PNG: "PNG",
        ImageFormat.JPEG: "JPEG",
        ImageFormat.WEBP: "WEBP",
        ImageFormat.BMP: "BMP",
        ImageFormat.TIFF: "TIFF",
        ImageFormat.GIF: "GIF",
    }

    EXTENSION_MAP = {
        ImageFormat.PNG: ".png",
        ImageFormat.JPEG: ".jpg",
        ImageFormat.WEBP: ".webp",
        ImageFormat.BMP: ".bmp",
        ImageFormat.TIFF: ".tiff",
        ImageFormat.GIF: ".gif",
    }

    MIME_MAP = {
        ".png": ImageFormat.PNG,
        ".jpg": ImageFormat.JPEG,
        ".jpeg": ImageFormat.JPEG,
        ".webp": ImageFormat.WEBP,
        ".bmp": ImageFormat.BMP,
        ".tiff": ImageFormat.TIFF,
        ".tif": ImageFormat.TIFF,
        ".gif": ImageFormat.GIF,
    }

    @classmethod
    def convert(
        cls,
        input_path: str,
        output_path: str,
        target_format: ImageFormat,
        quality: int = 85,
        progress_callback=None,
    ) -> str:
        if progress_callback:
            progress_callback(10)

        img = Image.open(input_path)

        if progress_callback:
            progress_callback(30)

        if img.mode == "RGBA" and target_format in (ImageFormat.JPEG, ImageFormat.BMP):
            background = Image.new("RGB", img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background
        elif img.mode == "P" and target_format != ImageFormat.GIF:
            img = img.convert("RGBA" if target_format == ImageFormat.PNG else "RGB")
        elif img.mode not in ("RGB", "RGBA", "L"):
            img = img.convert("RGB")

        if progress_callback:
            progress_callback(60)

        save_kwargs = {}
        pil_format = cls.FORMAT_MAP[target_format]

        if target_format in (ImageFormat.JPEG, ImageFormat.WEBP):
            save_kwargs["quality"] = quality
        if target_format == ImageFormat.JPEG:
            save_kwargs["optimize"] = True
        if target_format == ImageFormat.PNG:
            save_kwargs["optimize"] = True
        if target_format == ImageFormat.WEBP:
            save_kwargs["method"] = 4
        if target_format == ImageFormat.TIFF:
            save_kwargs["compression"] = "tiff_lzw"

        img.save(output_path, format=pil_format, **save_kwargs)

        if progress_callback:
            progress_callback(100)

        img.close()
        return output_path

    @classmethod
    def get_format_from_extension(cls, filename: str) -> ImageFormat | None:
        ext = Path(filename).suffix.lower()
        return cls.MIME_MAP.get(ext)

    @classmethod
    def get_extension(cls, fmt: ImageFormat) -> str:
        return cls.EXTENSION_MAP[fmt]

    @classmethod
    def validate_image(cls, file_path: str) -> dict:
        try:
            img = Image.open(file_path)
            info = {
                "valid": True,
                "width": img.size[0],
                "height": img.size[1],
                "mode": img.mode,
                "format": img.format,
            }
            img.close()
            return info
        except Exception as e:
            return {"valid": False, "error": str(e)}


class CSVConverter:
    @staticmethod
    def to_json(
        input_path: str,
        output_path: str,
        progress_callback=None,
    ) -> str:
        if progress_callback:
            progress_callback(10)

        with open(input_path, "r", encoding="utf-8-sig") as f:
            content = f.read()

        if progress_callback:
            progress_callback(20)

        dialect = None
        try:
            sample = content[:8192]
            dialect = csv.Sniffer().sniff(sample)
        except csv.Error:
            pass

        if progress_callback:
            progress_callback(30)

        reader_kwargs = {}
        if dialect:
            reader_kwargs["dialect"] = dialect

        reader = csv.DictReader(io.StringIO(content), **reader_kwargs)
        rows = []
        field_names = reader.fieldnames or []

        total_lines = content.count("\n")
        if total_lines == 0:
            total_lines = 1

        for i, row in enumerate(reader):
            cleaned_row = {}
            for key, value in row.items():
                if key is None:
                    continue
                clean_key = key.strip()
                if value is not None:
                    stripped = value.strip()
                    try:
                        if "." in stripped:
                            cleaned_row[clean_key] = float(stripped)
                        else:
                            cleaned_row[clean_key] = int(stripped)
                    except (ValueError, TypeError):
                        if stripped.lower() in ("true", "false"):
                            cleaned_row[clean_key] = stripped.lower() == "true"
                        elif stripped.lower() in ("null", "none", ""):
                            cleaned_row[clean_key] = None
                        else:
                            cleaned_row[clean_key] = stripped
                else:
                    cleaned_row[clean_key] = None
            rows.append(cleaned_row)

            if progress_callback and i % 100 == 0:
                progress = min(30 + int((i / total_lines) * 60), 90)
                progress_callback(progress)

        if progress_callback:
            progress_callback(90)

        result = {
            "metadata": {
                "total_records": len(rows),
                "fields": [f.strip() for f in field_names if f],
                "source_file": Path(input_path).name,
            },
            "data": rows,
        }

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2, ensure_ascii=False, default=str)

        if progress_callback:
            progress_callback(100)

        return output_path


class ConverterFactory:
    _converters = {
        "pdf_to_docx": PDFConverter,
        "image_convert": ImageConverter,
        "csv_to_json": CSVConverter,
    }

    @classmethod
    def get_converter(cls, conversion_type: str):
        converter_class = cls._converters.get(conversion_type)
        if not converter_class:
            raise ValueError(f"Unsupported conversion type: {conversion_type}")
        return converter_class

    @classmethod
    def supported_types(cls) -> list[str]:
        return list(cls._converters.keys())

    @classmethod
    def validate_input_format(cls, conversion_type: str, filename: str) -> bool:
        ext = Path(filename).suffix.lower()

        validation_map = {
            "pdf_to_docx": [".pdf"],
            "image_convert": [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif", ".gif"],
            "csv_to_json": [".csv"],
        }

        allowed = validation_map.get(conversion_type, [])
        return ext in allowed
